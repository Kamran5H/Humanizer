# -*- coding: utf-8 -*-
"""Generate a ~500-word chemistry test document (.docx) for the paraphraser."""
import sys
import docx

TITLE = "Failure Modes in Rechargeable Zinc-Air Batteries"

PARAS = [
    "Rechargeable zinc-air batteries are attractive for grid storage because they "
    "combine a high theoretical energy density of 1086 Wh/kg with low-cost, abundant "
    "materials. During discharge, the zinc anode is oxidized to zincate, Zn(OH)4^2-, "
    "which decomposes to ZnO once the electrolyte becomes locally saturated. At the air "
    "cathode, the oxygen reduction reaction consumes O2 and produces OH-, sustaining the "
    "alkaline environment. On charge, the oxygen evolution reaction regenerates O2 at a "
    "substantial overpotential, and this asymmetry limits the round-trip efficiency to "
    "roughly 60% in most reported cells [1].",

    "The air cathode is the principal source of loss. A practical electrode must host a "
    "bifunctional catalyst that is active for both the oxygen reduction reaction and the "
    "oxygen evolution reaction, yet few materials tolerate the oxidizing potentials near "
    "2.0 V without corrosion. Carbon supports, though conductive and porous, are "
    "themselves oxidized above 1.65 V, which degrades the gas diffusion layer and floods "
    "the triple-phase boundary. Researchers have therefore explored perovskites such as "
    "LaNiO3, spinels like Co3O4, and nitrogen-doped carbons as more stable alternatives [2].",

    "The zinc anode introduces a second set of problems. Repeated stripping and plating "
    "redistributes the active material, producing dendrites that can pierce the separator "
    "and short the cell. Passivation by a dense ZnO film raises the internal resistance "
    "and lowers the accessible specific capacity, often below 400 mAh/g in practice, far "
    "from the theoretical 820 mAh/g. Additives such as bismuth oxide and organic "
    "surfactants suppress dendrite growth, but they frequently reduce the coulombic "
    "efficiency at the same time [3].",

    "Carbonation is a slower, insidious failure mode. Atmospheric CO2 diffuses through the "
    "cathode and reacts with the KOH electrolyte to form K2CO3, which precipitates in the "
    "pores, blocks oxygen transport, and steadily raises the overpotential. Because the "
    "reaction is thermodynamically favorable, sealing and CO2 scrubbing only delay it "
    "rather than prevent it. Reported cells lose a measurable fraction of their capacity "
    "within a few hundred cycles once carbonation begins.",

    "Taken together, these mechanisms are coupled rather than independent. A catalyst "
    "chosen for a low oxygen evolution reaction overpotential may accelerate carbon "
    "corrosion; an additive that protects the zinc anode may poison the catalyst; "
    "increasing the current density to raise power density worsens both dendrite formation "
    "and local heating. This coupling explains why single-parameter optimization rarely "
    "translates into longer cycle life, and why the field still lacks a common reporting "
    "standard that would let independent groups compare round-trip efficiency, energy "
    "density, and cycle life on equal terms [4].",
]


def main(dst):
    d = docx.Document()
    d.add_heading(TITLE, level=1)
    for p in PARAS:
        d.add_paragraph(p)
    d.save(dst)
    wc = sum(len(p.split()) for p in PARAS)
    print("WORDS:", wc, "PARAS:", len(PARAS), "SAVED:", dst)


if __name__ == "__main__":
    main(sys.argv[1] if len(sys.argv) > 1 else "test_chem_500.docx")
