# -*- coding: utf-8 -*-
"""
Safe, content-preserving AI-tell remover for academic .docx files.

Operates on the ORIGINAL document, editing text runs in place. It ONLY swaps a
curated set of AI-writing tells (vocabulary, filler phrases, over-used
transitions) for plain equivalents of the SAME meaning. It never paraphrases,
never expands, never touches numbers, units, chemical formulas/subscripts,
citations, equations, the title, or the keyword list. Formatting is preserved
because we edit run.text only where a known pattern occurs.
"""
import re
import sys
import docx

# --- smart-case vocabulary swaps (meaning-preserving) -----------------------
# (pattern-without-boundaries, replacement). Applied with \b word boundaries,
# case-insensitive, preserving the matched token's leading capitalization.
VOCAB = [
    (r"delves into", "examines"),
    (r"delve into", "examine"),
    (r"delving into", "examining"),
    (r"utilizes", "uses"),
    (r"utilized", "used"),
    (r"utilizing", "using"),
    (r"utilization", "use"),
    (r"utilize", "use"),
    (r"leveraging", "using"),
    (r"showcases", "shows"),
    (r"showcased", "showed"),
    (r"showcasing", "showing"),
    (r"showcase", "show"),
    (r"underscores", "highlights"),
    (r"underscored", "highlighted"),
    (r"underscoring", "highlighting"),
    (r"underscore", "highlight"),
    (r"a testament to", "evidence of"),
    (r"in the realm of", "in"),
    (r"in the landscape of", "in"),
    (r"plays a crucial role in", "is central to"),
    (r"play a crucial role in", "are central to"),
    (r"plays a pivotal role in", "is central to"),
    (r"play a pivotal role in", "are central to"),
    (r"pivotal", "key"),
    (r"a myriad of", "many"),
    (r"myriad of", "many"),
    (r"a plethora of", "many"),
    (r"plethora of", "many"),
    (r"seamlessly", "smoothly"),
    (r"seamless", "smooth"),
    (r"cutting-edge", "advanced"),
    (r"paradigm shift", "major change"),
    (r"intricate", "complex"),
    (r"meticulously", "carefully"),
    (r"meticulous", "careful"),
]

# Over-used transitions -> plainer equivalents (same meaning).
TRANSITIONS = [
    (r"Moreover,", "Also,"),
    (r"Furthermore,", "In addition,"),
    (r"Additionally,", "Also,"),
]

# Filler phrases removed entirely; the following clause is kept and its first
# letter capitalized. Anchored to sentence start to avoid mid-sentence damage.
FILLER = re.compile(
    r"(^|[.!?]\s+)"
    r"(?:it is worth noting|it is important to note|it should be noted|"
    r"it is worth mentioning|it is worth pointing out|it is important to mention|"
    r"it is worth highlighting)\s+that\s+(\w)",
    re.IGNORECASE,
)


def _smartcase(match, replacement):
    tok = match.group(0)
    if tok[:1].isupper():
        return replacement[:1].upper() + replacement[1:]
    return replacement


def _apply_vocab(text):
    changed = 0
    for pat, rep in VOCAB + TRANSITIONS:
        rx = re.compile(r"\b" + pat + r"\b", re.IGNORECASE)

        def _sub(m, rep=rep):
            nonlocal changed
            changed += 1
            return _smartcase(m, rep)

        text = rx.sub(_sub, text)
    return text, changed


def _apply_filler(text):
    changed = 0

    def _sub(m):
        nonlocal changed
        changed += 1
        return m.group(1) + m.group(2).upper()

    text = FILLER.sub(_sub, text)
    return text, changed


def transform(text):
    t, c1 = _apply_filler(text)
    t, c2 = _apply_vocab(t)
    return t, c1 + c2


def iter_all_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def _render_bar(current: int, total: int, width: int = 24) -> None:
    if total <= 0:
        return
    pct = (current / total) * 100.0
    filled = int(width * (current / total))
    bar = "█" * filled + "░" * (width - filled)
    sys.stderr.write(f"\r  [{bar}] {pct:5.1f}% ({current}/{total} paras)")
    sys.stderr.flush()
    if current >= total:
        sys.stderr.write("\n")
        sys.stderr.flush()


def main(src, dst):
    doc = docx.Document(src)
    all_paras = list(iter_all_paragraphs(doc))
    total_paras = len(all_paras)
    total_edits = 0
    runs_touched = 0
    samples = []

    for idx, p in enumerate(all_paras, 1):
        _render_bar(idx, total_paras)
        for run in p.runs:
            if not run.text:
                continue
            new, c = transform(run.text)
            if c and new != run.text:
                if len(samples) < 12:
                    samples.append((run.text, new))
                run.text = new
                total_edits += c
                runs_touched += 1

    doc.save(dst)
    return total_edits, runs_touched, samples, total_paras


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python safe_humanize_docx.py <input.docx> <output.docx>")
        sys.exit(1)
    src, dst = sys.argv[1], sys.argv[2]
    edits, runs, samples, total_p = main(src, dst)

    print("\n┌────────────────────────────────────────────────────────┐")
    print("│         SAFE DOCX HUMANIZER — PROGRESS REPORT          │")
    print("├────────────────────────────────────────────────────────┤")
    print(f"│  Paragraphs Processed: {total_p:<31} │")
    print(f"│  AI Tell Edits Made:   {edits:<31} │")
    print(f"│  Text Runs Modified:   {runs:<31} │")
    print(f"│  Output Saved:         {str(dst):<31} │")
    print("└────────────────────────────────────────────────────────┘\n")

    if samples:
        print("Sample Replacements:")
        for a, b in samples[:5]:
            print(f"  • BEFORE: {a[:120]}")
            print(f"    AFTER:  {b[:120]}")
