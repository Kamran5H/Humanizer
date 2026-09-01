# -*- coding: utf-8 -*-
"""Generate a ~1000-word, deliberately AI-styled chemistry document (.docx) with
two graphs, one schematic figure, captions/descriptions, and references +
inline citations. Intended as a realistic 'looks AI-written' input for the
constrained paraphraser."""
import os
import sys
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyArrow, Rectangle
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_DIR = os.path.dirname(os.path.abspath(__file__))


def make_schematic(path):
    fig, ax = plt.subplots(figsize=(6, 3.2))
    ax.add_patch(Rectangle((0.05, 0.2), 0.18, 0.6, fc="#9fb6d1", ec="k"))
    ax.add_patch(Rectangle((0.23, 0.2), 0.44, 0.6, fc="#cfe6cf", ec="k"))
    ax.add_patch(Rectangle((0.67, 0.2), 0.18, 0.6, fc="#e6cfcf", ec="k"))
    ax.text(0.14, 0.5, "Zn\nanode", ha="center", va="center", fontsize=10)
    ax.text(0.45, 0.5, "KOH\nelectrolyte", ha="center", va="center", fontsize=10)
    ax.text(0.76, 0.5, "air\ncathode", ha="center", va="center", fontsize=10)
    ax.add_patch(FancyArrow(0.86, 0.85, 0.10, 0.0, width=0.008,
                            length_includes_head=True, color="k"))
    ax.text(0.90, 0.90, "O2 in", fontsize=9)
    ax.add_patch(FancyArrow(0.23, 0.9, -0.10, 0.0, width=0.008,
                            length_includes_head=True, color="k"))
    ax.text(0.02, 0.92, "e- (discharge)", fontsize=9)
    ax.set_xlim(0, 1.05)
    ax.set_ylim(0, 1.05)
    ax.axis("off")
    fig.tight_layout()
    fig.savefig(path, dpi=150)
    plt.close(fig)


def make_efficiency_graph(path):
    cycles = list(range(0, 210, 10))
    fig, ax = plt.subplots(figsize=(6, 3.4))
    for label, e0, decay in [("Co3O4", 63, 0.055), ("LaNiO3", 60, 0.038),
                             ("N-doped C", 58, 0.085)]:
        y = [e0 * (2.718 ** (-decay * (c / 20.0))) + 20 for c in cycles]
        ax.plot(cycles, y, marker="o", ms=3, label=label)
    ax.set_xlabel("Cycle number")
    ax.set_ylabel("Round-trip efficiency (%)")
    ax.set_ylim(20, 70)
    ax.legend(frameon=False, fontsize=9)
    ax.grid(alpha=0.3)
    fig.tight_layout()
    fig.savefig(path, dpi=150)
    plt.close(fig)


def make_polarization_graph(path):
    j = [i / 10.0 for i in range(0, 201, 5)]  # mA/cm2
    dis = [1.45 - 0.0016 * x - 0.02 * (x ** 0.5) for x in j]
    chg = [1.95 + 0.0020 * x + 0.03 * (x ** 0.5) for x in j]
    fig, ax = plt.subplots(figsize=(6, 3.4))
    ax.plot(j, dis, label="discharge", color="#2a6f2a")
    ax.plot(j, chg, label="charge", color="#a33")
    ax.fill_between(j, dis, chg, color="#ddd", alpha=0.4)
    ax.set_xlabel("Current density (mA/cm2)")
    ax.set_ylabel("Cell voltage (V)")
    ax.legend(frameon=False, fontsize=9)
    ax.grid(alpha=0.3)
    fig.tight_layout()
    fig.savefig(path, dpi=150)
    plt.close(fig)


ABSTRACT = (
    "Rechargeable zinc-air batteries have emerged as a compelling candidate for "
    "next-generation grid storage, leveraging a myriad of favorable attributes that "
    "collectively underscore their transformative potential. In this work we delve into "
    "the intricate interplay of electrochemical mechanisms that govern their performance, "
    "and we highlight the pivotal role that bifunctional electrocatalysis plays in "
    "shaping round-trip efficiency. Moreover, we critically examine the degradation "
    "pathways that continue to hinder practical deployment."
)

SECTIONS = [
    ("Introduction",
     "The global transition toward renewable energy has spurred an unprecedented demand "
     "for low-cost, sustainable storage solutions. In the realm of metal-air chemistries, "
     "the zinc-air battery stands out as a particularly promising platform, owing to its "
     "high theoretical energy density of 1086 Wh/kg, the natural abundance of zinc, and "
     "its inherent safety in aqueous alkaline media [1]. Furthermore, unlike lithium-based "
     "systems, zinc-air cells operate without the flammable organic electrolytes that pose "
     "significant safety concerns. It is worth noting that these advantages have driven a "
     "surge of research interest over the past decade, positioning the technology as a "
     "cornerstone of future decarbonization efforts [2]. Nevertheless, a considerable gulf "
     "persists between the theoretical promise of the chemistry and the modest performance "
     "of laboratory prototypes, a discrepancy that motivates the present review. In the "
     "following sections we systematically unpack the mechanistic origins of this gap and "
     "delineate the material and engineering levers that may ultimately close it."),
    ("Operating Principles",
     "During discharge, the zinc anode undergoes oxidation to form zincate, Zn(OH)4^2-, "
     "which subsequently decomposes into ZnO once the local electrolyte becomes saturated. "
     "At the air cathode, the oxygen reduction reaction consumes O2 and generates OH-, "
     "thereby sustaining the alkaline environment that is crucial for continued operation. "
     "Upon charging, the oxygen evolution reaction regenerates O2 at a substantial "
     "overpotential, and this pronounced asymmetry between the two processes plays a "
     "pivotal role in limiting the round-trip efficiency to roughly 60% in most reported "
     "cells [3]. Figure 1 provides a schematic overview of the cell architecture and the "
     "principal charge-transfer pathways. The thermodynamic equilibrium potential of the "
     "oxygen electrode sits at 1.23 V versus the reversible hydrogen electrode, yet the "
     "sluggish kinetics of both reactions impose kinetic overpotentials that dwarf this "
     "figure in practice. As a result, the voltaic efficiency of the cell is dictated less "
     "by thermodynamics than by the catalytic activity of the electrode surface."),
    ("Bifunctional Electrocatalysis",
     "The air cathode represents the principal source of energy loss and, consequently, "
     "the primary target for materials innovation. A practical electrode must host a "
     "bifunctional catalyst that remains active toward both the oxygen reduction reaction "
     "and the oxygen evolution reaction, yet remarkably few materials can tolerate the "
     "oxidizing potentials near 2.0 V without succumbing to corrosion. Carbon supports, "
     "though highly conductive and porous, are themselves oxidized above 1.65 V, which "
     "degrades the gas diffusion layer and floods the triple-phase boundary. Researchers "
     "have therefore explored a diverse array of alternatives, including perovskites such "
     "as LaNiO3, spinels like Co3O4, and nitrogen-doped carbons [4]. As shown in Figure 2, "
     "the round-trip efficiency of these catalysts decays at markedly different rates over "
     "extended cycling, underscoring the importance of catalyst stability. Increasingly, "
     "attention has turned to atomically dispersed transition-metal sites anchored on "
     "nitrogen-doped carbon, which offer a myriad of tunable coordination environments and "
     "unusually high mass activity. Nonetheless, translating the impressive performance of "
     "such catalysts from the rotating disk electrode to a practical air cathode remains a "
     "formidable and largely unsolved challenge."),
    ("Electrolyte and Additive Strategies",
     "The electrolyte is far more than a passive ionic conductor; it plays a decisive role "
     "in dictating both zinc utilization and cathode longevity. Concentrated KOH solutions "
     "afford excellent conductivity, but they simultaneously accelerate zinc corrosion and "
     "aggravate carbonation. To mitigate these effects, researchers have investigated a "
     "broad palette of additives, including potassium fluoride, polyethylene glycol, and "
     "various ionic liquids, each of which modifies the zinc deposition morphology in "
     "subtle ways. Gel and quasi-solid electrolytes have likewise attracted considerable "
     "attention, since they suppress water loss and dendrite penetration while preserving "
     "acceptable ionic mobility. Collectively, these strategies illustrate the delicate "
     "balance that must be struck between conductivity, stability, and manufacturability."),
    ("Degradation Pathways",
     "Beyond the cathode, the zinc anode introduces a second constellation of challenges. "
     "Repeated stripping and plating redistributes the active material, producing dendrites "
     "that can pierce the separator and short the cell. Passivation by a dense ZnO film "
     "raises the internal resistance and lowers the accessible specific capacity, often "
     "below 400 mAh/g in practice, far from the theoretical 820 mAh/g [5]. Carbonation "
     "constitutes a slower yet insidious failure mode, whereby atmospheric CO2 reacts with "
     "the KOH electrolyte to form K2CO3, which precipitates within the pores and steadily "
     "raises the overpotential [6]. Figure 3 presents representative charge-discharge "
     "polarization curves, illustrating the voltage gap that widens as current density "
     "increases. Crucially, these degradation modes rarely act in isolation; instead, they "
     "reinforce one another in a vicious cycle, whereby dendrite-induced shorting locally "
     "heats the electrolyte, which in turn accelerates both carbonation and the parasitic "
     "hydrogen evolution reaction. Disentangling the relative contribution of each pathway "
     "therefore demands carefully designed operando diagnostics rather than post-mortem "
     "analysis alone."),
    ("Outlook",
     "Taken together, these mechanisms are deeply coupled rather than independent. A "
     "catalyst selected for a low oxygen evolution reaction overpotential may inadvertently "
     "accelerate carbon corrosion, while an additive that protects the zinc anode may "
     "poison the catalyst. This intricate coupling explains why single-parameter "
     "optimization rarely translates into meaningful gains in cycle life [7]. Ultimately, "
     "the field would benefit enormously from a common reporting standard that enables "
     "independent groups to compare round-trip efficiency, energy density, and cycle life "
     "on truly equal terms [8]. Looking ahead, the convergence of machine-learning-guided "
     "catalyst discovery, advanced operando spectroscopy, and rationally engineered "
     "electrolytes offers a genuinely exciting pathway forward. By embracing this "
     "integrated, systems-level perspective, the community stands poised to unlock the "
     "full and long-anticipated promise of rechargeable zinc-air technology."),
]

REFERENCES = [
    "Lee, J. et al. Metal-air batteries for grid storage. Nat. Energy 5, 230-245 (2020).",
    "Wang, H. & Xu, Q. Progress in zinc-air chemistry. Chem. Rev. 121, 4562-4610 (2021).",
    "Fu, J. et al. Electrically rechargeable zinc-air batteries. Adv. Mater. 29, 1604685 (2017).",
    "Gu, P. et al. Bifunctional oxygen electrocatalysts. Energy Environ. Sci. 11, 2925 (2018).",
    "Mainar, A. R. et al. Zinc anode challenges. J. Energy Storage 15, 304-328 (2018).",
    "Schroder, D. et al. Carbonation in alkaline cells. J. Power Sources 266, 302 (2014).",
    "Pan, J. et al. Coupled degradation in zinc-air cells. ACS Energy Lett. 4, 1112 (2019).",
    "Li, Y. & Dai, H. Reporting standards for metal-air batteries. Chem. Soc. Rev. 43, 5257 (2014).",
]

CAPTIONS = [
    ("Figure 1.", "Schematic of a rechargeable zinc-air cell showing the zinc anode, the "
     "alkaline KOH electrolyte, and the porous air cathode, together with the oxygen "
     "inlet and the direction of electron flow during discharge."),
    ("Figure 2.", "Round-trip efficiency as a function of cycle number for three "
     "bifunctional catalysts (Co3O4, LaNiO3, and nitrogen-doped carbon), highlighting "
     "the divergent stability of the candidate materials over 200 cycles."),
    ("Figure 3.", "Charge and discharge polarization curves plotted against current "
     "density; the shaded region denotes the voltage gap that governs the round-trip "
     "efficiency of the cell."),
]


def _cap(doc, tag, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{tag} ")
    r.bold = True
    r2 = p.add_run(text)
    r2.italic = True
    r2.font.size = Pt(9)


def main(dst):
    s1 = os.path.join(OUT_DIR, "_fig1_schematic.png")
    s2 = os.path.join(OUT_DIR, "_fig2_efficiency.png")
    s3 = os.path.join(OUT_DIR, "_fig3_polarization.png")
    make_schematic(s1)
    make_efficiency_graph(s2)
    make_polarization_graph(s3)

    d = docx.Document()
    d.add_heading("Coupled Failure Modes and Bounded Efficiency in "
                  "Rechargeable Zinc-Air Batteries", level=0)
    d.add_heading("Abstract", level=2)
    d.add_paragraph(ABSTRACT)

    figs = [s1, s2, s3]
    body_words = len(ABSTRACT.split())
    for i, (title, body) in enumerate(SECTIONS):
        d.add_heading(title, level=1)
        d.add_paragraph(body)
        body_words += len(body.split())
        # place a figure after sections that reference one (Principles, Cat, Degrad)
        if title == "Operating Principles":
            d.add_picture(s1, width=Inches(4.5)); _last_center(d)
            _cap(d, *CAPTIONS[0])
        elif title == "Bifunctional Electrocatalysis":
            d.add_picture(s2, width=Inches(4.8)); _last_center(d)
            _cap(d, *CAPTIONS[1])
        elif title == "Degradation Pathways":
            d.add_picture(s3, width=Inches(4.8)); _last_center(d)
            _cap(d, *CAPTIONS[2])

    d.add_heading("References", level=1)
    for i, r in enumerate(REFERENCES, 1):
        d.add_paragraph(f"[{i}] {r}")

    d.save(dst)
    for f in figs:
        try:
            os.remove(f)
        except OSError:
            pass
    print("WORDS(body+abstract):", body_words, "SAVED:", dst)


def _last_center(doc):
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


if __name__ == "__main__":
    main(sys.argv[1] if len(sys.argv) > 1 else "test_chem_1000_AI.docx")
