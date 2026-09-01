"""
Humanizer GUI Module.
Editorial Tkinter application with custom SleekProgress, HiDPI support,
live AI score gauge badge, and responsive multithreaded processing.
"""

from __future__ import annotations

import logging
import os
import random
import sys
import threading
from pathlib import Path
from tkinter import filedialog, messagebox, ttk
import tkinter as tk
from typing import Optional

from humanizer.config import HumanizerConfig
from humanizer.detector import score_text
from humanizer.engine import humanize_text, humanize_docx
from humanizer.rules import ensure_nltk_data

logger = logging.getLogger(__name__)


class ResultDialog(tk.Toplevel):
    """Modern comprehensive executive progress and audit report dialog."""
    def __init__(
        self,
        parent,
        final_score: float,
        initial_score: float = 0.0,
        path: Optional[Path] = None,
        metrics: Optional[dict] = None,
        duration: float = 0.0,
        word_count: int = 0,
        para_count: int = 0,
    ) -> None:
        super().__init__(parent)
        self.title("Humanization Progress & Audit Report")
        self.geometry("540x440")
        self.resizable(False, False)
        self.configure(bg="#F5F1E8")
        self.transient(parent)
        self.grab_set()

        frm = ttk.Frame(self, padding=20)
        frm.pack(fill="both", expand=True)

        # Header Badge
        is_passing = final_score <= 5.0
        color = "#4A6D3F" if is_passing else ("#B8892B" if final_score <= 20.0 else "#8B2A1F")
        verdict = "100% Human Passing — Verified Safe" if is_passing else "Humanized"

        tk.Label(frm, text="Executive Audit Report", font=("Georgia", 16, "bold"), fg="#1A1613", bg="#F5F1E8").pack(anchor="w", pady=(0, 2))
        tk.Label(frm, text=verdict, font=("Segoe UI", 10, "italic", "bold"), fg=color, bg="#F5F1E8").pack(anchor="w", pady=(0, 12))

        # Metrics Card Frame
        card = tk.Frame(frm, bg="#FFFFFF", highlightbackground="#D8D0BA", highlightthickness=1, padx=16, pady=12)
        card.pack(fill="both", expand=True, pady=(0, 14))

        # Row 1: AI Scores
        delta = initial_score - final_score
        delta_txt = f" (↓ {delta:.1f}% reduction)" if delta > 0 else ""
        self._add_stat_row(card, "Initial AI Score:", f"{initial_score:.1f}%", row=0)
        self._add_stat_row(card, "Final AI Score:", f"{final_score:.1f}%{delta_txt}", row=1, val_color=color, val_bold=True)

        # Row 2: Content volume
        self._add_stat_row(card, "Paragraphs Processed:", f"{para_count} paragraphs", row=2)
        self._add_stat_row(card, "Total Word Count:", f"{word_count:,} words", row=3)

        # Row 3: Detection Signals
        if metrics:
            self._add_stat_row(card, "Sentence Burstiness (CV):", f"{metrics.get('burstiness', 0.0):.1f}", row=4)
            self._add_stat_row(card, "Token Perplexity:", f"{metrics.get('perplexity', 0.0):.1f}", row=5)
            self._add_stat_row(card, "Residual AI Tells:", f"{metrics.get('ai_tells', 0)} detected", row=6)

        if duration > 0:
            self._add_stat_row(card, "Processing Time:", f"{duration:.1f}s", row=7)

        if path:
            lbl_p = tk.Label(frm, text=f"Saved to: {path.name}", font=("Segoe UI", 9), fg="#8B8474", bg="#F5F1E8")
            lbl_p.pack(anchor="w", pady=(0, 10))

        # Button Bar
        btn_bar = tk.Frame(frm, bg="#F5F1E8")
        btn_bar.pack(fill="x", pady=(4, 0))

        btn_copy = ttk.Button(btn_bar, text="Copy Report", style="Ghost.TButton", command=lambda: self._copy_report(parent, initial_score, final_score, word_count, para_count, metrics, duration))
        btn_copy.pack(side="left")

        btn_done = ttk.Button(btn_bar, text="Done", style="Accent.TButton", command=self.destroy)
        btn_done.pack(side="right")

    def _add_stat_row(self, parent: tk.Frame, label: str, val: str, row: int, val_color: str = "#1A1613", val_bold: bool = False) -> None:
        font_val = ("Segoe UI", 10, "bold") if val_bold else ("Segoe UI", 10)
        tk.Label(parent, text=label, font=("Segoe UI", 10), fg="#6B655A", bg="#FFFFFF").grid(row=row, column=0, sticky="w", pady=3)
        tk.Label(parent, text=val, font=font_val, fg=val_color, bg="#FFFFFF").grid(row=row, column=1, sticky="e", padx=(20, 0), pady=3)
        parent.grid_columnconfigure(0, weight=1)

    def _copy_report(self, parent, init_sc, final_sc, words, paras, metrics, duration) -> None:
        rep = (
            f"=== HUMANIZER PRO AUDIT REPORT ===\n"
            f"Verdict:           {'100% Human Passing' if final_sc <= 5.0 else 'Humanized'}\n"
            f"Initial AI Score:  {init_sc:.1f}%\n"
            f"Final AI Score:    {final_sc:.1f}% (Reduction: {init_sc - final_sc:.1f}%)\n"
            f"Paragraphs:        {paras}\n"
            f"Words:             {words}\n"
            f"Burstiness:        {metrics.get('burstiness', 0.0) if metrics else 0.0}\n"
            f"Perplexity:        {metrics.get('perplexity', 0.0) if metrics else 0.0}\n"
            f"AI Tells:          {metrics.get('ai_tells', 0) if metrics else 0}\n"
            f"Time:              {duration:.1f}s\n"
        )
        parent.clipboard_clear()
        parent.clipboard_append(rep)
        messagebox.showinfo("Report Copied", "Progress and audit report copied to clipboard.")


class HumanizerApp(tk.Tk):
    _loaded_path: Optional[Path] = None

    PAL = {
        "bg": "#F5F1E8",
        "bg2": "#EEE9DC",
        "surface": "#FFFFFF",
        "surface2": "#FBF7EE",
        "shadow": "#D6CFBE",
        "rule": "#D8D0BA",
        "text": "#1A1613",
        "muted": "#6B655A",
        "dim": "#8B8474",
        "accent": "#8B6F1F",
        "accent2": "#B8892B",
        "green": "#4A6D3F",
        "red": "#8B2A1F",
    }

    FONT_DISPLAY = ("Georgia", 24, "bold")
    FONT_SUB = ("Georgia", 11, "italic")
    FONT_BODY = ("Georgia", 11)
    FONT_UI = ("Segoe UI", 10)
    FONT_BTN_HERO = ("Georgia", 12, "bold")
    FONT_BTN = ("Segoe UI", 9)

    def __init__(self) -> None:
        super().__init__()
        self._enable_hidpi()
        self.title("Humanizer Pro — Academic & Scientific Stealth Suite")
        self.geometry("1240x820")
        self.minsize(920, 620)
        self.configure(bg=self.PAL["bg"])

        self._init_style()
        self._cancel_event = threading.Event()
        self._start_time = 0.0
        self._initial_score = 0.0
        self._build_ui()
        threading.Thread(target=ensure_nltk_data, daemon=True).start()

    def _enable_hidpi(self) -> None:
        if sys.platform == "win32":
            try:
                from ctypes import windll
                windll.shcore.SetProcessDpiAwareness(2)
            except Exception:
                pass

    def _init_style(self) -> None:
        p = self.PAL
        style = ttk.Style()
        style.theme_use("clam")
        style.configure(".", background=p["bg"], foreground=p["text"])
        style.configure("Accent.TButton", background=p["accent"], foreground="#FFFFFF", font=self.FONT_BTN_HERO, padding=(32, 10))
        style.map("Accent.TButton", background=[("active", p["accent2"])])
        style.configure("Ghost.TButton", background=p["bg"], foreground=p["muted"], font=self.FONT_BTN, padding=(12, 6))

        # Progress bar styling
        style.configure(
            "Custom.Horizontal.TProgressbar",
            troughcolor=p["bg2"],
            background=p["accent"],
            darkcolor=p["accent"],
            lightcolor=p["accent2"],
            bordercolor=p["rule"],
            thickness=8,
        )

    def _build_ui(self) -> None:
        # Header
        hdr = tk.Frame(self, bg=self.PAL["bg"], padx=24, pady=16)
        hdr.pack(fill="x")

        tk.Label(hdr, text="Humanizer Pro", font=self.FONT_DISPLAY, fg=self.PAL["text"], bg=self.PAL["bg"]).pack(anchor="w")
        tk.Label(hdr, text="Academic & Journal-Grade Stealth Rewriting Engine", font=self.FONT_SUB, fg=self.PAL["muted"], bg=self.PAL["bg"]).pack(anchor="w")

        # Controls Toolbar
        tb = tk.Frame(self, bg=self.PAL["bg"], padx=24, pady=8)
        tb.pack(fill="x")

        self.btn_run = ttk.Button(tb, text="Humanize", style="Accent.TButton", command=self._start_humanize)
        self.btn_run.pack(side="left", padx=(0, 12))

        self.btn_cancel = ttk.Button(tb, text="Cancel", style="Ghost.TButton", command=self._cancel, state="disabled")
        self.btn_cancel.pack(side="left", padx=(0, 12))

        self.btn_load = ttk.Button(tb, text="Open File...", style="Ghost.TButton", command=self._open_file)
        self.btn_load.pack(side="left", padx=(0, 8))

        self.btn_save = ttk.Button(tb, text="Save Output...", style="Ghost.TButton", command=self._save_file)
        self.btn_save.pack(side="left", padx=(0, 8))

        self.btn_score = ttk.Button(tb, text="Check Score", style="Ghost.TButton", command=self._score_input)
        self.btn_score.pack(side="left", padx=(0, 8))

        self.lbl_score_badge = tk.Label(tb, text="AI Score: —", font=("Segoe UI", 10, "bold"), fg=self.PAL["muted"], bg=self.PAL["bg"])
        self.lbl_score_badge.pack(side="right", padx=8)

        # Panes
        panes = tk.Frame(self, bg=self.PAL["bg"], padx=24, pady=8)
        panes.pack(fill="both", expand=True)
        panes.columnconfigure(0, weight=1)
        panes.columnconfigure(1, weight=1)
        panes.rowconfigure(0, weight=1)

        # Left: Original
        f_left = tk.Frame(panes, bg=self.PAL["surface"], highlightbackground=self.PAL["rule"], highlightthickness=1)
        f_left.grid(row=0, column=0, sticky="nsew", padx=(0, 8))
        self.txt_in = tk.Text(f_left, font=self.FONT_BODY, wrap="word", bd=0, padx=12, pady=12)
        self.txt_in.pack(fill="both", expand=True)

        # Right: Humanized
        f_right = tk.Frame(panes, bg=self.PAL["surface"], highlightbackground=self.PAL["rule"], highlightthickness=1)
        f_right.grid(row=0, column=1, sticky="nsew", padx=(8, 0))
        self.txt_out = tk.Text(f_right, font=self.FONT_BODY, wrap="word", bd=0, padx=12, pady=12)
        self.txt_out.pack(fill="both", expand=True)

        # Progress Frame (Progress Bar + Percentage Badge)
        p_frame = tk.Frame(self, bg=self.PAL["bg2"], padx=16, pady=4)
        p_frame.pack(fill="x", side="bottom")

        self.progress_bar = ttk.Progressbar(
            p_frame,
            orient="horizontal",
            mode="determinate",
            maximum=100.0,
            style="Custom.Horizontal.TProgressbar",
        )
        self.progress_bar.pack(fill="x", side="left", expand=True, padx=(0, 12), pady=2)

        self.lbl_progress_pct = tk.Label(
            p_frame,
            text="0.0%",
            font=("Segoe UI", 9, "bold"),
            fg=self.PAL["accent"],
            bg=self.PAL["bg2"],
            width=6,
            anchor="e",
        )
        self.lbl_progress_pct.pack(side="right")

        # Status Bar
        self.lbl_status = tk.Label(
            self,
            text="Ready",
            font=("Segoe UI", 9),
            fg=self.PAL["muted"],
            bg=self.PAL["bg2"],
            anchor="w",
            padx=16,
            pady=2,
        )
        self.lbl_status.pack(fill="x", side="bottom")

    def _open_file(self) -> None:
        p = filedialog.askopenfilename(filetypes=[("Documents", "*.docx *.txt *.pdf")])
        if not p:
            return
        path = Path(p)
        self._loaded_path = path
        if path.suffix.lower() == ".docx":
            import docx
            doc = docx.Document(str(path))
            text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
            self.txt_in.delete("1.0", "end")
            self.txt_in.insert("1.0", text)
            self.lbl_status.config(text=f"Loaded Word Document: {path.name}")
        else:
            text = path.read_text(encoding="utf-8", errors="replace")
            self.txt_in.delete("1.0", "end")
            self.txt_in.insert("1.0", text)
            self.lbl_status.config(text=f"Loaded text file: {path.name}")
        self._score_input()

    def _save_file(self) -> None:
        text = self.txt_out.get("1.0", "end-1c")
        if not text.strip():
            return
        p = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Text file", "*.txt"), ("Word Document", "*.docx")])
        if not p:
            return
        target = Path(p)
        if target.suffix.lower() == ".docx":
            import docx
            doc = docx.Document()
            for para in text.split("\n\n"):
                if para.strip():
                    doc.add_paragraph(para.strip())
            doc.save(str(target))
        else:
            target.write_text(text, encoding="utf-8")
        self.lbl_status.config(text=f"Saved to: {target.name}")

    def _score_input(self) -> None:
        text = self.txt_out.get("1.0", "end-1c").strip() or self.txt_in.get("1.0", "end-1c").strip()
        if not text:
            return
        res = score_text(text)
        sc = res["ai_score"]
        color = self.PAL["green"] if sc <= 5.0 else (self.PAL["accent"] if sc <= 20.0 else self.PAL["red"])
        self.lbl_score_badge.config(text=f"AI Score: {sc:.1f}%", fg=color)

    def _update_progress(self, current: int, total: int, pct: float, message: str) -> None:
        self.progress_bar["value"] = pct
        self.lbl_progress_pct.config(text=f"{pct:5.1f}%")
        self.lbl_status.config(text=f"{message} ({current}/{total})" if total > 0 else message)

    def _start_humanize(self) -> None:
        text = self.txt_in.get("1.0", "end-1c")
        if not text.strip():
            return

        import time
        self._start_time = time.time()
        init_res = score_text(text)
        self._initial_score = init_res["ai_score"]

        self._cancel_event.clear()
        self.btn_run.config(state="disabled")
        self.btn_cancel.config(state="normal")
        self.progress_bar["value"] = 0.0
        self.lbl_progress_pct.config(text="0.0%")
        self.lbl_status.config(text="Starting stealth humanization...")

        cfg = HumanizerConfig(
            use_gemini=True,
            stealth_level=3,
            target_pct=5.0,
            scientific=True,
            cancel_event=self._cancel_event,
            status_callback=lambda s: self.after(0, self.lbl_status.config, {"text": s}),
            progress_callback=lambda cur, tot, pct, msg: self.after(0, self._update_progress, cur, tot, pct, msg),
        )

        def _bg():
            try:
                if self._loaded_path and self._loaded_path.suffix.lower() == ".docx":
                    saved = humanize_docx(self._loaded_path, cfg)
                    import docx
                    doc = docx.Document(str(saved))
                    out_text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
                else:
                    out_text = humanize_text(text, cfg)
                    saved = None

                self.after(0, self._finish_humanize, out_text, saved)
            except Exception as e:
                self.after(0, self._error_humanize, str(e))

        threading.Thread(target=_bg, daemon=True).start()

    def _cancel(self) -> None:
        self._cancel_event.set()
        self.lbl_status.config(text="Cancelling...")

    def _finish_humanize(self, out_text: str, saved_path: Optional[Path]) -> None:
        import time
        duration = time.time() - self._start_time
        self.txt_out.delete("1.0", "end")
        self.txt_out.insert("1.0", out_text)
        self.btn_run.config(state="normal")
        self.btn_cancel.config(state="disabled")
        self.progress_bar["value"] = 100.0
        self.lbl_progress_pct.config(text="100.0%")
        self._score_input()

        res = score_text(out_text)
        words = len(out_text.split())
        paras = len([p for p in out_text.split("\n\n") if p.strip()])

        ResultDialog(
            self,
            final_score=res["ai_score"],
            initial_score=self._initial_score,
            path=saved_path,
            metrics=res,
            duration=duration,
            word_count=words,
            para_count=paras,
        )

    def _error_humanize(self, err_msg: str) -> None:
        self.btn_run.config(state="normal")
        self.btn_cancel.config(state="disabled")
        self.progress_bar["value"] = 0.0
        self.lbl_progress_pct.config(text="0.0%")
        self.lbl_status.config(text=f"Error: {err_msg}")
        messagebox.showerror("Humanizer Error", err_msg)


def launch_gui() -> None:
    app = HumanizerApp()
    app.mainloop()
