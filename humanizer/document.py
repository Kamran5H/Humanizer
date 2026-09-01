"""
High-Fidelity Document Processing Module.
Supports .docx with run-level style preservation (bold, italics, citations, sub/superscript, colors),
along with PDF extraction, TXT handling, and persistent session checkpointing.
"""

from __future__ import annotations

import json
import logging
import os
import re
import threading
import time
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Callable, Generator, Optional

import docx
from docx import Document
from docx.shared import RGBColor

from humanizer.config import HumanizerConfig

logger = logging.getLogger(__name__)

_SKIP_STYLES = {"code", "verbatim", "source code", "header", "footer"}
REF_LINE = re.compile(r"^\s*\[\d+\]")
CAPTION_LINE = re.compile(r"^\s*(Figure|Fig\.|Table)\s*\d", re.IGNORECASE)


class Checkpoint:
    """Per-session persistent state for a humanize run."""
    DIR = Path(os.environ.get("LOCALAPPDATA") or Path.home() / "AppData/Local") / "humanizer_pro" / "sessions"

    def __init__(
        self,
        sid: str,
        input_text: str,
        cfg_key: dict,
        total: int,
        mode: str = "text",
        source_path: Optional[str] = None,
    ) -> None:
        self.sid = sid
        self.path = self.DIR / f"{sid}.json"
        self.total = total
        self.mode = mode
        self.source_path = source_path
        self._blocks: dict[str, str] = {}
        self._save_lock = threading.Lock()
        self._created_at = datetime.now().isoformat(timespec="seconds")
        self._load_existing(input_text, cfg_key)
        self._input = input_text
        self._cfg_key = cfg_key

    def _load_existing(self, input_text: str, cfg_key: dict) -> None:
        if not self.path.exists():
            return
        try:
            d = json.loads(self.path.read_text(encoding="utf-8"))
            if d.get("input") == input_text and d.get("cfg") == cfg_key:
                self._blocks = dict(d.get("blocks") or {})
                self._created_at = d.get("created_at", self._created_at)
        except Exception:
            pass

    def get(self, idx: int) -> Optional[str]:
        return self._blocks.get(str(idx))

    def done_count(self) -> int:
        return len(self._blocks)

    def clear(self) -> None:
        try:
            if self.path.exists():
                self.path.unlink()
        except Exception:
            pass

    def put(self, idx: int, text: str) -> None:
        with self._save_lock:
            self._blocks[str(idx)] = text
            try:
                self.DIR.mkdir(parents=True, exist_ok=True)
                payload = {
                    "version": 2,
                    "sid": self.sid,
                    "mode": self.mode,
                    "source_path": self.source_path,
                    "input": self._input,
                    "cfg": self._cfg_key,
                    "total": self.total,
                    "blocks": self._blocks,
                    "created_at": self._created_at,
                    "updated_at": datetime.now().isoformat(timespec="seconds"),
                }
                tmp = self.path.with_suffix(".tmp")
                tmp.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
                tmp.replace(self.path)
            except Exception as e:
                logger.warning(f"Checkpoint save failed: {e}")

    @staticmethod
    def compute_id(input_text: str, cfg_key: dict) -> str:
        import hashlib
        h = hashlib.sha1()
        h.update(input_text.encode("utf-8", errors="replace"))
        h.update(json.dumps(cfg_key, sort_keys=True).encode("utf-8"))
        return h.hexdigest()[:16]


@dataclass
class RunStyle:
    bold: Optional[bool] = None
    italic: Optional[bool] = None
    underline: Optional[bool] = None
    subscript: Optional[bool] = None
    superscript: Optional[bool] = None
    font_name: Optional[str] = None
    font_size: Optional[object] = None
    font_color: Optional[RGBColor] = None
    style_name: Optional[str] = None


def _extract_run_style(run) -> RunStyle:
    rs = RunStyle(
        bold=run.bold,
        italic=run.italic,
        underline=run.underline,
        subscript=run.font.subscript if run.font else None,
        superscript=run.font.superscript if run.font else None,
    )
    if run.font:
        rs.font_name = run.font.name
        rs.font_size = run.font.size
        if run.font.color and run.font.color.rgb:
            rs.font_color = run.font.color.rgb
    if run.style:
        try:
            rs.style_name = run.style.name
        except Exception:
            pass
    return rs


def _apply_run_style(run, rs: RunStyle) -> None:
    if rs.bold is not None:
        run.bold = rs.bold
    if rs.italic is not None:
        run.italic = rs.italic
    if rs.underline is not None:
        run.underline = rs.underline
    if run.font:
        if rs.subscript is not None:
            run.font.subscript = rs.subscript
        if rs.superscript is not None:
            run.font.superscript = rs.superscript
        if rs.font_name:
            run.font.name = rs.font_name
        if rs.font_size:
            run.font.size = rs.font_size
        if rs.font_color:
            run.font.color.rgb = rs.font_color
    if rs.style_name:
        try:
            run.style = rs.style_name
        except Exception:
            pass


def _tag_runs(para) -> tuple[str, dict[str, dict]]:
    """Encodes styled runs as semantic tags and protected placeholders."""
    tagged_parts = []
    vault: dict[str, dict] = {}
    counter = 0

    for run in para.runs:
        t = run.text
        if not t:
            continue
        xml_str = run._element.xml
        has_complex = (
            "<w:drawing" in xml_str
            or "<w:pict" in xml_str
            or "<m:oMath" in xml_str
            or "<w:object" in xml_str
        )
        is_cite = bool(run.font and (run.font.superscript or run.font.subscript) or re.fullmatch(r"\[[a-zA-Z0-9,\-\s]{1,15}\]", t.strip()))

        if has_complex or is_cite:
            tag = f"__RUN_LOCKED_{counter}__"
            vault[tag] = {"text": t, "style": _extract_run_style(run), "locked": True}
            tagged_parts.append(tag)
            counter += 1
        elif run.bold and run.italic:
            tag_open = f"<bi id='{counter}'>"
            tag_close = f"</bi id='{counter}'>"
            vault[f"bi_{counter}"] = {"style": _extract_run_style(run), "locked": False}
            tagged_parts.append(f"{tag_open}{t}{tag_close}")
            counter += 1
        elif run.bold:
            tag_open = f"<b id='{counter}'>"
            tag_close = f"</b id='{counter}'>"
            vault[f"b_{counter}"] = {"style": _extract_run_style(run), "locked": False}
            tagged_parts.append(f"{tag_open}{t}{tag_close}")
            counter += 1
        elif run.italic:
            tag_open = f"<i id='{counter}'>"
            tag_close = f"</i id='{counter}'>"
            vault[f"i_{counter}"] = {"style": _extract_run_style(run), "locked": False}
            tagged_parts.append(f"{tag_open}{t}{tag_close}")
            counter += 1
        else:
            tagged_parts.append(t)

    return "".join(tagged_parts), vault


def _reconstruct_runs(para, rewritten_text: str, vault: dict[str, dict], base_style: RunStyle) -> None:
    """Clears and rebuilds paragraph runs faithfully preserving all inline styles."""
    # First restore locked placeholders
    locked_tokens = [k for k, v in vault.items() if v.get("locked")]
    if locked_tokens:
        for tok in locked_tokens:
            val = vault[tok]["text"]
            rewritten_text = rewritten_text.replace(tok, val)

    # Parse semantic tags: <b id='X'>text</b> or <i id='Y'>text</i> or <bi id='Z'>text</bi>
    tag_re = re.compile(r"<(b|i|bi)\s+id=['\"](\d+)['\"]>(.*?)</\1(?:\s+id=['\"]\2['\"])?>", re.DOTALL | re.IGNORECASE)
    
    segments = []
    last_idx = 0
    for m in tag_re.finditer(rewritten_text):
        if m.start() > last_idx:
            plain_chunk = rewritten_text[last_idx:m.start()]
            if plain_chunk:
                segments.append((plain_chunk, base_style))
        tag_type = m.group(1).lower()
        tag_id = m.group(2)
        tag_content = m.group(3)
        key = f"{tag_type}_{tag_id}"
        st = vault.get(key, {}).get("style", base_style)
        segments.append((tag_content, st))
        last_idx = m.end()

    if last_idx < len(rewritten_text):
        trailing = rewritten_text[last_idx:]
        if trailing:
            segments.append((trailing, base_style))

    # Strip any stray unclosed tags
    clean_segments = []
    for s_txt, s_style in segments:
        clean_txt = re.sub(r"</?[bi|b|i]\b[^>]*>", "", s_txt)
        if clean_txt:
            clean_segments.append((clean_txt, s_style))

    # If no segments, write plain
    if not clean_segments:
        clean_segments = [(re.sub(r"</?[bi|b|i]\b[^>]*>", "", rewritten_text), base_style)]

    # Clear existing runs in paragraph
    for r in para.runs:
        r.text = ""
    # Add new runs with precise styling
    for seg_text, st in clean_segments:
        if not seg_text:
            continue
        new_run = para.add_run(seg_text)
        _apply_run_style(new_run, st)


def should_skip_paragraph(para) -> bool:
    t = para.text.strip()
    if not t:
        return True
    style = (para.style.name or "").lower() if para.style else ""
    if any(s in style for s in _SKIP_STYLES):
        return True
    if REF_LINE.match(t) or CAPTION_LINE.match(t):
        return True
    return False


def iter_document_paragraphs(doc: Document) -> Generator:
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
