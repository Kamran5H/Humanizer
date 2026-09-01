"""
Humanizer Headless CLI Interface.
Supports file processing (.docx, .txt), stdin pipelines, score reporting,
live console percentage progress bar, and executive audit reports.
"""

from __future__ import annotations

import argparse
import json
import sys
import time
from pathlib import Path

from humanizer.config import HumanizerConfig
from humanizer.detector import score_text
from humanizer.engine import humanize_text, humanize_docx
from humanizer.lint import lint_text


def render_cli_progress_bar(current: int, total: int, pct: float, message: str, width: int = 24) -> None:
    """Render a live ASCII percentage progress bar directly to stderr."""
    if total <= 0:
        return
    filled = int(width * (pct / 100.0))
    filled = max(0, min(width, filled))
    bar = "█" * filled + "░" * (width - filled)
    msg_clean = message.replace("\n", " ")[:32]
    sys.stderr.write(f"\r  [{bar}] {pct:5.1f}% ({current}/{total}) | {msg_clean:<32}")
    sys.stderr.flush()
    if current >= total:
        sys.stderr.write("\n")
        sys.stderr.flush()


def print_cli_report(
    initial_score: float,
    final_score: float,
    words: int,
    paras: int,
    metrics: dict | None = None,
    dst: Path | None = None,
    duration: float = 0.0,
) -> None:
    """Print a clean executive progress and audit report box."""
    is_passing = final_score <= 5.0
    verdict = "100% Human Passing (Verified Safe)" if is_passing else "Humanized"
    delta = initial_score - final_score
    delta_str = f" (↓ {delta:.1f}% reduction)" if delta > 0 else ""

    lines = [
        "┌────────────────────────────────────────────────────────────────────────┐",
        "│                HUMANIZER PRO — PROGRESS & AUDIT REPORT                 │",
        "├────────────────────────────────────────────────────────────────────────┤",
        f"│  Verdict:             {verdict:<47} │",
        f"│  Initial AI Score:    {initial_score:>5.1f}%{'':<43} │",
        f"│  Final AI Score:      {final_score:>5.1f}%{delta_str:<43} │",
        f"│  Paragraphs Processed: {paras:<47} │",
        f"│  Total Word Count:    {words:<47} │",
    ]

    if metrics:
        lines.append(f"│  Burstiness (CV):     {metrics.get('burstiness', 0.0):>5.1f}{'':<43} │")
        lines.append(f"│  Perplexity:          {metrics.get('perplexity', 0.0):>5.1f}{'':<43} │")
        lines.append(f"│  Residual AI Tells:   {metrics.get('ai_tells', 0):<47} │")

    if duration > 0:
        lines.append(f"│  Processing Time:     {duration:>5.1f}s{'':<43} │")

    if dst:
        dst_name = dst.name if len(str(dst.name)) <= 46 else dst.name[:43] + "..."
        lines.append(f"│  Output Saved:        {dst_name:<47} │")

    lines.append("└────────────────────────────────────────────────────────────────────────┘")

    print("\n" + "\n".join(lines) + "\n", file=sys.stderr)


def build_parser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(
        prog="humanizer_pro",
        description="Humanizer Pro — High-Performance Stealth Text Humanizer & Academic Rewriter",
    )
    p.add_argument("file", nargs="?", help="Input file to humanize (.docx or .txt)")
    p.add_argument("--text", help="Raw input string to humanize")
    p.add_argument("--stdin", action="store_true", help="Read input from stdin")
    p.add_argument("--out", "-o", help="Explicit output path")
    p.add_argument("--level", type=int, choices=[1, 2, 3], default=3, help="Stealth level (1=fast, 2=gauge, 3=stealth+refine)")
    p.add_argument("--target", type=float, default=5.0, help="Target AI score threshold (default: 5.0)")
    p.add_argument("--budget", type=int, default=6, help="Retry budget per paragraph (default: 6)")
    p.add_argument("--scientific", action="store_true", default=True, help="Enable academic/scientific register")
    p.add_argument("--general", dest="scientific", action="store_false", help="Use general/casual register")
    p.add_argument("--local", action="store_true", help="Force offline rule pipeline (no cloud calls)")
    p.add_argument("--score", action="store_true", help="Print AI detection score breakdown")
    p.add_argument("--report", action="store_true", default=True, help="Display executive audit report box")
    p.add_argument("--no-report", dest="report", action="store_false", help="Suppress executive audit report box")
    p.add_argument("--json", action="store_true", help="Output results and metrics as structured JSON")
    p.add_argument("--workers", type=int, default=4, help="Parallel concurrency workers")
    return p


def main(argv: list[str] | None = None) -> int:
    parser = build_parser()
    args = parser.parse_args(argv)

    show_bar = not args.json and (args.file or args.text or not args.stdin)
    p_cb = render_cli_progress_bar if show_bar else None

    cfg = HumanizerConfig(
        use_gemini=not args.local,
        stealth_level=args.level,
        target_pct=args.target,
        retry_budget=args.budget,
        scientific=args.scientific,
        parallel_workers=args.workers,
        progress_callback=p_cb,
    )

    t0 = time.time()

    if args.stdin:
        input_text = sys.stdin.read()
        init_res = score_text(input_text)
        out_text = humanize_text(input_text, cfg)
        scores = score_text(out_text)
        dur = time.time() - t0

        if args.json:
            report = {
                "success": True,
                "output": out_text,
                "initial_score": init_res.get("ai_score", 0.0),
                "final_score": scores.get("ai_score", 0.0),
                "metrics": scores,
                "duration_seconds": round(dur, 2),
                "lint": lint_text(out_text, input_text),
            }
            print(json.dumps(report, indent=2))
        else:
            sys.stdout.write(out_text)
            if args.report or args.score:
                paras = len([p for p in out_text.split("\n\n") if p.strip()])
                words = len(out_text.split())
                print_cli_report(init_res["ai_score"], scores["ai_score"], words, paras, scores, duration=dur)
        return 0

    if args.text:
        init_res = score_text(args.text)
        out_text = humanize_text(args.text, cfg)
        scores = score_text(out_text)
        dur = time.time() - t0

        if args.json:
            report = {
                "success": True,
                "output": out_text,
                "initial_score": init_res.get("ai_score", 0.0),
                "final_score": scores.get("ai_score", 0.0),
                "metrics": scores,
                "duration_seconds": round(dur, 2),
                "lint": lint_text(out_text, args.text),
            }
            print(json.dumps(report, indent=2))
        else:
            print(out_text)
            if args.report or args.score:
                paras = len([p for p in out_text.split("\n\n") if p.strip()])
                words = len(out_text.split())
                print_cli_report(init_res["ai_score"], scores["ai_score"], words, paras, scores, duration=dur)
        return 0

    if args.file:
        src = Path(args.file)
        if not src.exists():
            print(f"Error: File '{src}' does not exist.", file=sys.stderr)
            return 1

        dst = Path(args.out) if args.out else src.with_name(f"{src.stem}_humanized{src.suffix}")

        if src.suffix.lower() == ".docx":
            import docx
            doc_in = docx.Document(str(src))
            in_text = "\n\n".join(p.text for p in doc_in.paragraphs if p.text.strip())
            init_res = score_text(in_text)

            saved = humanize_docx(src, cfg, dst=dst)

            doc_out = docx.Document(str(saved))
            full_text = "\n\n".join(p.text for p in doc_out.paragraphs if p.text.strip())
            scores = score_text(full_text)
        else:
            raw = src.read_text(encoding="utf-8")
            init_res = score_text(raw)
            out_text = humanize_text(raw, cfg)
            dst.write_text(out_text, encoding="utf-8")
            scores = score_text(out_text)
            full_text = out_text

        dur = time.time() - t0
        paras = len([p for p in full_text.split("\n\n") if p.strip()])
        words = len(full_text.split())

        if args.json:
            report = {
                "success": True,
                "input_file": str(src),
                "output_file": str(dst),
                "initial_score": init_res.get("ai_score", 0.0),
                "final_score": scores.get("ai_score", 0.0),
                "metrics": scores,
                "duration_seconds": round(dur, 2),
            }
            print(json.dumps(report, indent=2))
        else:
            if args.report or args.score:
                print_cli_report(init_res["ai_score"], scores["ai_score"], words, paras, scores, dst=dst, duration=dur)
            else:
                print(f"Saved: {dst}")
        return 0

    # No arguments provided -> launch GUI
    from humanizer.gui import launch_gui
    launch_gui()
    return 0


if __name__ == "__main__":
    sys.exit(main())
